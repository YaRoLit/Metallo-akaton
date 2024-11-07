import docx


class GOST_parser():
    def __init__(self, gost_fname: str, db_adder, llm_request):
        self.db_adder = db_adder
        self.llm_request = llm_request
        self.gost = docx.Document(gost_fname)
        self.gost_num = self.search_doc_num(self.gost)
        self.parse_table(self.gost)
        self.parse_text(self.gost)


    def search_doc_num(self, gost: docx.Document) -> None:
        """
        Определяем номер ГОСТа в формате [номер, год]
        """
        doc_head = ''
        for paragraph in gost.paragraphs[:20]:
            doc_head += paragraph.text.strip() + '\t'
        doc_head = self._rm_double_rows(doc_head)
        doc_head = self._rm_extra_spaces(doc_head)
        prompt = f"""Ты исследователь текстов, который абсолютно точно соблюдает инструкции.
Представлен титульный лист нормативно-технического документа:
{doc_head}
Определи, имеется ли в титульном листе идентификатор нормативно-технического документа (НТД).
Идентификатор НТД состоит из нескольких частей. Сначала идет название НТД (например, 'ГОСТ' или 'ТУ').
После этого идет номер НТД (например, '12345'). В конце идет год выпуска НТД (например, '2010' или '85').
Примеры обозначения НТД в тексте: ГОСТ 123456-2015, ГОСТ 222222 2022, ТУ 543-86.
В ответе выведи номер документа без дополнительных комментариев. 
Если номера нет, напиши просто 'нет': 
Примеры ответов:
ГОСТ 123456-2015
ГОСТ 222222-2022
ТУ 543-1986
нет"""     
        answer = self.llm_request(prompt)
        if answer != "нет":
            try:
                gost_num = answer.split(' ')
                gost_num = gost_num[1].split('-')
                return gost_num
            except:
                print("Ошибка парсинга номера НТД")
                return False


    def parse_text(self, gost: docx.Document) -> None:
        """
        Парсинг текста из docx документа
        """
        text = ''
        for paragraph in gost.paragraphs:
            if ("таблица" in paragraph.text.lower()[:10]) or ("рисунок" in paragraph.text.lower()[:10]):
                continue
            text += paragraph.text + '\n'
        text = self._rm_double_rows(text)
        text = self._rm_extra_spaces(text)
        prompt = f"""Ты исследователь текстов, который абсолютно точно соблюдает инструкции.
Имеется текст нормативно-технического документа:
{text}
Этот текст необходимо разбить на отдельные фрагменты.
Фрагменты должны быть как можно меньшего размера, но при этом не должен теряться их смысл. 
В выводе напечатай этот текст, разделенный на фрагменты.
Для разделения фрагменты используй строку '----'.
В выводе должен быть только текст без комментариев. Пример ответа:
----
3.23 высокий отпуск: Технологический процесс нагрева проката ниже температуры , выдержки и охлаждения его с заданной скоростью или на спокойном воздухе.
----
3.24 механическое старение: Процесс искусственного старения в соответствии с ГОСТ 7268.
----"""
        answer = self.llm_request(prompt)
        text = answer.split("----")
        for idx, block in enumerate(text):
            self.db_adder(
                doc_text = block,
                doc_meta = {
                    'gost_num': self.gost_num[0],
                    'gost_year': self.gost_num[1],
                    'type': 'text',                   
                    },
                    ids=f'gost{self.gost_num[0]}_block{idx}'
            )


    def parse_table(self, gost: docx.Document) -> None:
        """
        Парсинг таблиц из docx документа
        """
        table_name = False
        for block in gost.iter_inner_content():
            if (type(block) == docx.text.paragraph.Paragraph):
                if ("таблица" in block.text.lower()[:10]):
                    table_name = self._rm_extra_spaces(block.text).lower()
            elif type(block) == docx.table.Table:
                if not table_name:
                    continue
                table_text = ''
                table_num = table_name.split()[1].upper()
                for row in block.rows:
                    table_text += '\n'
                    row_text = ''
                    for cell in row.cells:
                        #if cell.grid_span > 1:
                        #    continue
                        row_text = row_text + cell.text + '\t'
                    table_text += row_text
                table_text = self._rm_double_rows(table_text)
                table_meta = self.analyse_table(table_name, table_text)
                self.db_adder(
                    doc_text = table_meta,
                    doc_meta = {
                        'gost_num': self.gost_num[0],
                        'gost_year': self.gost_num[1],
                        'type': 'table_meta',
                        'table_name': table_name,
                        'table_num': table_num                        
                    },
                    ids=f'gost{self.gost_num[0]}_table_{table_num}_meta'
                )
                table_name = False


    def analyse_table(self, table_name, table_body: str)  -> str:
        """
        Анализируем таблицу llm и выводим ее метаданные
        """
        prompt = f"""Ты исследователь текстов, который абсолютно точно соблюдает инструкции.
Представлена следующая таблица:
{table_name}
{table_body}
Укажи, какая информация содержится в таблице (в соответствии с наименованиями столбцов и строк).
Также перечисли все уникальные элементы, содержащиеся в конкретном столбце
Если несколько столбцов объединены одним заголовком, нужно указать заголовок, а в качестве его уникальных элементов названия этих столбцов.
Например, если столбцы ."Предел текучести", "Временное сопротивление", "Относительное удлинение" объединены заголовком "Механические свойства стали",
нужно указать в выводе: "Механические свойства стали": ["Предел текучести", "Временное сопротивление", "Относительное удлинение"].
Далее нужно перечислить эти столбцы и их уникальные элементы.
Примечания к таблице в ответе указывать не нужно.
В ответе первой строкой укажи название таблицы.
Ответ выведи в виде json. Не пиши никаких дополнительных комментариев и надписей кроме блока json в скобках {{}}.
Примеры ответа:
{{
"Название таблицы": "Таблица 1 - наименование таблицы"
"Класс прочности": ["135", "250", "325"],
"Марка стали": ["А2В", "ГСС2Н", "Ф14"],
"Механические свойства стали": ["Предел текучести", "Временное сопротивление", "Относительное удлинение"],
"Предел текучести": ["385", "400", "450"],
"Временное сопротивление": ["300", "400", "500"]
"Относительное удлинение": ["19", "20", "22"]
}}
{{
"Название таблицы": "Таблица 4 - содержание таблицы"
"Ударная вязкость, Д ж /см 2, не менее, при тем пературе испы тания, °С": ["-40", "0", "+40"],
"Химические элементы": ["C", "S", "N"],
}}"""
        answer = self.llm_request(prompt)
        return answer


    @staticmethod
    def _rm_extra_spaces(text: str) -> str:
        """
        Убираем лишние переносы строк
        """
        text = text.replace(u'\xa0', u' ')
        text = text.replace('  ', ' ')
        text = text.replace('   ', ' ')
        text = text.strip()
        return text


    @staticmethod
    def _rm_double_rows(text: str) -> str:
        """
        Убираем повторящиеся строки
        """
        text = text.replace("\n\n", "\n")
        text = text.replace("\n\n\n", "\n")
        text = text.replace("\n\n", "\n")
        text = text.split('\n')
        new_text = [text[0]]
        for line in text:
            if new_text[-1].strip() == line.strip():
                continue
            else:
                new_text.append(line)
        return '\n'.join(new_text)
